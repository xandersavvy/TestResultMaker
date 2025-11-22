"""
Test Evidence Helper - Desktop Application
A tool for manual software testers to capture screenshots and notes,
organize them by test steps, and export to Word documents.

Requirements:
    pip install pyqt5 keyboard python-docx pandas pillow openpyxl pyperclip
"""

import sys
import os
import threading
import tempfile
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Optional, Dict, Any
from dataclasses import dataclass, field
import re
import platform

# Third-party imports
try:
    from PyQt5.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QPushButton, QLabel, QTextEdit, QFileDialog, QLineEdit, 
        QMessageBox, QStatusBar, QGroupBox, QTableWidget, QTableWidgetItem,
        QDialog, QDialogButtonBox, QHeaderView, QComboBox, QCheckBox
    )
    from PyQt5.QtCore import Qt, QTimer, pyqtSignal, QObject
    from PyQt5.QtGui import QIcon, QFont
except ImportError as e:
    print(f"PyQt5 not found: {e}. Install via: pip install pyqt5")
    sys.exit(1)

try:
    import keyboard
except ImportError as e:
    print(f"keyboard module not found: {e}. Install via: pip install keyboard")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError as e:
    print(f"python-docx not found: {e}. Install via: pip install python-docx")
    sys.exit(1)

try:
    from PIL import ImageGrab, Image
except ImportError as e:
    print(f"PIL/Pillow not found: {e}. Install via: pip install pillow")
    sys.exit(1)

try:
    import pandas as pd
except ImportError as e:
    print(f"pandas not found: {e}. Install via: pip install pandas openpyxl")
    sys.exit(1)

try:
    import pyperclip
except ImportError as e:
    print(f"pyperclip not found: {e}. Install via: pip install pyperclip")
    sys.exit(1)


# ==================== CONFIGURATION ====================
CAPTURE_HOTKEY = "ctrl+alt+s"
NEXT_STEP_HOTKEY = "ctrl+alt+right"
PREV_STEP_HOTKEY = "ctrl+alt+left"

# Logging configuration
LOG_DIR = Path.home() / ".test_evidence_helper"
LOG_DIR.mkdir(exist_ok=True)
LOG_FILE = LOG_DIR / "app.log"

# ==================== LOGGING SETUP ====================
def setup_logging() -> logging.Logger:
    """Configure application logging with rotation and proper formatting."""
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.DEBUG)
    
    # Avoid duplicate handlers
    if logger.handlers:
        return logger
    
    # File handler with rotation
    from logging.handlers import RotatingFileHandler
    file_handler = RotatingFileHandler(
        LOG_FILE, 
        maxBytes=5*1024*1024,  # 5MB
        backupCount=3
    )
    file_handler.setLevel(logging.DEBUG)
    
    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    
    # Formatter
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(funcName)s:%(lineno)d - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)
    
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    
    return logger


logger = setup_logging()


# ==================== DATA MODELS ====================
@dataclass
class Step:
    """Represents a single test step with evidence."""
    index: int
    title: str = ""
    expected: str = ""
    actual: str = ""
    notes: List[str] = field(default_factory=list)
    screenshots: List[str] = field(default_factory=list)  # file paths

    def add_note(self, text: str) -> None:
        """Add a note to this step if not empty."""
        if text and text.strip():
            self.notes.append(text.strip())
            logger.debug(f"Added note to step {self.index}: {text[:50]}...")

    def add_screenshot(self, path: str) -> bool:
        """Add a screenshot if the file exists."""
        if path and os.path.exists(path):
            self.screenshots.append(path)
            logger.info(f"Added screenshot to step {self.index}: {path}")
            return True
        logger.warning(f"Screenshot path invalid or doesn't exist: {path}")
        return False

    def is_empty(self) -> bool:
        """Check if step has any content."""
        return not any([
            self.title.strip(),
            self.expected.strip(),
            self.actual.strip(),
            self.notes,
            self.screenshots
        ])


# ==================== CLIPBOARD PARSER ====================
class ClipboardParser:
    """Parse clipboard content for Excel data (TSV/CSV format like Xray)."""
    
    @staticmethod
    def parse_excel_from_clipboard() -> Optional[pd.DataFrame]:
        """
        Parse Excel data from clipboard (tab-separated or comma-separated).
        Supports formats from Excel, Google Sheets, Jira Xray, etc.
        
        Returns:
            DataFrame if successful, None otherwise
        """
        try:
            # Try to get text from clipboard
            clipboard_text = pyperclip.paste()
            
            if not clipboard_text or not clipboard_text.strip():
                logger.warning("Clipboard is empty")
                return None
            
            # Try tab-separated first (most common from Excel/Xray)
            if '\t' in clipboard_text:
                logger.info("Detected tab-separated data in clipboard")
                from io import StringIO
                df = pd.read_csv(StringIO(clipboard_text), sep='\t', encoding='utf-8')
                return df
            
            # Try comma-separated
            elif ',' in clipboard_text:
                logger.info("Detected comma-separated data in clipboard")
                from io import StringIO
                df = pd.read_csv(StringIO(clipboard_text), sep=',', encoding='utf-8')
                return df
            
            # Try line-by-line (simple list)
            else:
                lines = [line.strip() for line in clipboard_text.split('\n') if line.strip()]
                if len(lines) > 0:
                    logger.info("Detected line-separated data in clipboard")
                    # Create single-column dataframe
                    df = pd.DataFrame({'Step': lines})
                    return df
            
            logger.warning("Could not parse clipboard as tabular data")
            return None
            
        except Exception as e:
            logger.error(f"Failed to parse clipboard: {e}", exc_info=True)
            return None
    
    @staticmethod
    def detect_columns(df: pd.DataFrame) -> Dict[str, Optional[str]]:
        """
        Auto-detect column mappings for test steps.
        
        Args:
            df: Input dataframe
            
        Returns:
            Dictionary with detected column names
        """
        columns_lower = [str(c).lower().strip() for c in df.columns]
        mapping = {
            'title': None,
            'expected': None,
            'actual': None,
            'description': None
        }
        
        # Detect title/step column
        title_keywords = ['step', 'title', 'test', 'case', 'action', 'description', 'name', 'summary']
        for idx, col_name in enumerate(columns_lower):
            if any(keyword in col_name for keyword in title_keywords):
                mapping['title'] = df.columns[idx]
                break
        
        # Default to first column if not found
        if mapping['title'] is None and len(df.columns) > 0:
            mapping['title'] = df.columns[0]
        
        # Detect expected result column
        expected_keywords = ['expected', 'result', 'outcome', 'should']
        for idx, col_name in enumerate(columns_lower):
            if any(keyword in col_name for keyword in expected_keywords):
                mapping['expected'] = df.columns[idx]
                break
        
        # Default to second column if not found and exists
        if mapping['expected'] is None and len(df.columns) > 1:
            mapping['expected'] = df.columns[1]
        
        # Detect actual result column
        actual_keywords = ['actual', 'status', 'pass', 'fail']
        for idx, col_name in enumerate(columns_lower):
            if any(keyword in col_name for keyword in actual_keywords):
                mapping['actual'] = df.columns[idx]
                break
        
        logger.info(f"Detected column mapping: {mapping}")
        return mapping


# ==================== COLUMN MAPPING DIALOG ====================
class ColumnMappingDialog(QDialog):
    """Dialog for mapping Excel columns to step fields."""
    
    def __init__(self, df: pd.DataFrame, auto_mapping: Dict[str, Optional[str]], parent=None):
        super().__init__(parent)
        
        self.df = df
        self.mapping = auto_mapping.copy()
        
        self.setWindowTitle("Map Excel Columns")
        self.setMinimumWidth(500)
        
        layout = QVBoxLayout()
        
        # Info label
        info = QLabel(
            "Map your Excel columns to test step fields.\n"
            "Auto-detected mappings are pre-selected."
        )
        info.setWordWrap(True)
        layout.addWidget(info)
        
        # Preview table
        preview_group = QGroupBox("Data Preview (first 5 rows)")
        preview_layout = QVBoxLayout()
        
        self.preview_table = QTableWidget()
        self.preview_table.setRowCount(min(5, len(df)))
        self.preview_table.setColumnCount(len(df.columns))
        self.preview_table.setHorizontalHeaderLabels([str(c) for c in df.columns])
        
        for row_idx in range(min(5, len(df))):
            for col_idx, col_name in enumerate(df.columns):
                value = str(df.iloc[row_idx, col_idx])
                if pd.isna(df.iloc[row_idx, col_idx]):
                    value = ""
                item = QTableWidgetItem(value)
                item.setFlags(item.flags() & ~Qt.ItemIsEditable)
                self.preview_table.setItem(row_idx, col_idx, item)
        
        self.preview_table.resizeColumnsToContents()
        preview_layout.addWidget(self.preview_table)
        preview_group.setLayout(preview_layout)
        layout.addWidget(preview_group)
        
        # Mapping controls
        mapping_group = QGroupBox("Column Mappings")
        mapping_layout = QVBoxLayout()
        
        self.combos = {}
        column_options = ["<None>"] + [str(c) for c in df.columns]
        
        # Title mapping
        title_layout = QHBoxLayout()
        title_layout.addWidget(QLabel("Step Title/Description:"))
        self.combos['title'] = QComboBox()
        self.combos['title'].addItems(column_options)
        if auto_mapping['title']:
            self.combos['title'].setCurrentText(str(auto_mapping['title']))
        title_layout.addWidget(self.combos['title'])
        mapping_layout.addLayout(title_layout)
        
        # Expected mapping
        expected_layout = QHBoxLayout()
        expected_layout.addWidget(QLabel("Expected Result:"))
        self.combos['expected'] = QComboBox()
        self.combos['expected'].addItems(column_options)
        if auto_mapping['expected']:
            self.combos['expected'].setCurrentText(str(auto_mapping['expected']))
        expected_layout.addWidget(self.combos['expected'])
        mapping_layout.addLayout(expected_layout)
        
        # Actual mapping (optional)
        actual_layout = QHBoxLayout()
        actual_layout.addWidget(QLabel("Actual Result (optional):"))
        self.combos['actual'] = QComboBox()
        self.combos['actual'].addItems(column_options)
        if auto_mapping.get('actual'):
            self.combos['actual'].setCurrentText(str(auto_mapping['actual']))
        actual_layout.addWidget(self.combos['actual'])
        mapping_layout.addLayout(actual_layout)
        
        mapping_group.setLayout(mapping_layout)
        layout.addWidget(mapping_group)
        
        # Options
        options_group = QGroupBox("Import Options")
        options_layout = QVBoxLayout()
        
        self.skip_empty_checkbox = QCheckBox("Skip rows with empty title")
        self.skip_empty_checkbox.setChecked(True)
        options_layout.addWidget(self.skip_empty_checkbox)
        
        options_group.setLayout(options_layout)
        layout.addWidget(options_group)
        
        # Buttons
        buttons = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
        
        self.setLayout(layout)
    
    def get_mapping(self) -> Dict[str, Any]:
        """Get the final column mapping and options."""
        return {
            'title': self.combos['title'].currentText() if self.combos['title'].currentText() != "<None>" else None,
            'expected': self.combos['expected'].currentText() if self.combos['expected'].currentText() != "<None>" else None,
            'actual': self.combos['actual'].currentText() if self.combos['actual'].currentText() != "<None>" else None,
            'skip_empty': self.skip_empty_checkbox.isChecked()
        }


# ==================== HOTKEY MANAGER ====================
class HotkeyManager(QObject):
    """Manages global keyboard shortcuts in a separate thread."""
    
    capture_triggered = pyqtSignal()
    next_triggered = pyqtSignal()
    prev_triggered = pyqtSignal()
    error_occurred = pyqtSignal(str)
    
    def __init__(self, 
                 capture_key: str = CAPTURE_HOTKEY,
                 next_key: str = NEXT_STEP_HOTKEY,
                 prev_key: str = PREV_STEP_HOTKEY):
        super().__init__()
        self.capture_key = capture_key
        self.next_key = next_key
        self.prev_key = prev_key
        self._registered = False
        self._thread: Optional[threading.Thread] = None
        logger.info(f"HotkeyManager initialized with: capture={capture_key}, "
                   f"next={next_key}, prev={prev_key}")
    
    def register(self) -> bool:
        """Register all global hotkeys in a background thread."""
        if self._registered:
            logger.warning("Hotkeys already registered")
            return False
        
        try:
            keyboard.add_hotkey(self.capture_key, self._on_capture)
            keyboard.add_hotkey(self.next_key, self._on_next)
            keyboard.add_hotkey(self.prev_key, self._on_prev)
            self._registered = True
            
            # Start listener thread
            self._thread = threading.Thread(target=self._listen, daemon=True)
            self._thread.start()
            
            logger.info("Global hotkeys registered successfully")
            return True
        except Exception as e:
            error_msg = f"Failed to register hotkeys: {e}"
            logger.error(error_msg, exc_info=True)
            # self.error_occurred.emit(error_msg)
            return False
    
    def unregister(self) -> None:
        """Unregister all hotkeys."""
        if not self._registered:
            return
        
        try:
            keyboard.unhook_all_hotkeys()
            self._registered = False
            logger.info("Hotkeys unregistered")
        except Exception as e:
            logger.error(f"Error unregistering hotkeys: {e}", exc_info=True)
    
    def _listen(self) -> None:
        """Background thread listener (blocks forever)."""
        try:
            keyboard.wait()
        except Exception as e:
            logger.error(f"Hotkey listener thread error: {e}", exc_info=True)
    
    def _on_capture(self) -> None:
        """Capture hotkey callback."""
        try:
            self.capture_triggered.emit()
        except Exception as e:
            logger.error(f"Error in capture callback: {e}", exc_info=True)
    
    def _on_next(self) -> None:
        """Next step hotkey callback."""
        try:
            self.next_triggered.emit()
        except Exception as e:
            logger.error(f"Error in next callback: {e}", exc_info=True)
    
    def _on_prev(self) -> None:
        """Previous step hotkey callback."""
        try:
            self.prev_triggered.emit()
        except Exception as e:
            logger.error(f"Error in prev callback: {e}", exc_info=True)


# ==================== DOCX EXPORTER ====================
class DocxExporter:
    """Handles Word document generation from test steps."""
    
    @staticmethod
    def export(steps: List[Step], filepath: str, title: str = "Manual Test Execution") -> bool:
        """
        Export steps to a Word document.
        
        Args:
            steps: List of Step objects
            filepath: Output file path
            title: Document title
        
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document()
            
            # Title
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Metadata
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Total Steps: {len(steps)}")
            doc.add_page_break()
            
            # Steps
            for step in steps:
                if step.is_empty():
                    logger.debug(f"Skipping empty step {step.index}")
                    continue
                
                # Step heading
                doc.add_heading(f"Step {step.index}: {step.title or 'Untitled'}", level=1)
                
                # Expected vs Actual table
                if step.expected or step.actual:
                    table = doc.add_table(rows=2, cols=2)
                    table.style = 'Light Grid Accent 1'
                    
                    # Headers
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Expected Result"
                    hdr_cells[1].text = "Actual Result"
                    
                    # Content
                    content_cells = table.rows[1].cells
                    content_cells[0].text = step.expected or "N/A"
                    content_cells[1].text = step.actual or "N/A"
                    
                    doc.add_paragraph()  # spacing
                
                # Notes
                if step.notes:
                    doc.add_heading("Notes:", level=2)
                    for note in step.notes:
                        doc.add_paragraph(note, style='List Bullet')
                
                # Screenshots
                if step.screenshots:
                    doc.add_heading("Screenshots:", level=2)
                    for idx, screenshot_path in enumerate(step.screenshots, 1):
                        if not os.path.exists(screenshot_path):
                            logger.warning(f"Screenshot not found: {screenshot_path}")
                            doc.add_paragraph(f"[Screenshot {idx} missing: {screenshot_path}]")
                            continue
                        
                        try:
                            # Add caption
                            doc.add_paragraph(f"Screenshot {idx}:", style='Caption')
                            
                            # Add image with max width
                            doc.add_picture(screenshot_path, width=Inches(5.5))
                            
                            # Add spacing
                            doc.add_paragraph()
                        except Exception as img_error:
                            logger.error(f"Failed to add image {screenshot_path}: {img_error}")
                            doc.add_paragraph(f"[Error loading screenshot {idx}]")
                
                doc.add_page_break()
            
            # Save
            doc.save(filepath)
            logger.info(f"Document exported successfully to {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to export document: {e}", exc_info=True)
            return False


# ==================== MAIN WINDOW ====================
class MainWindow(QMainWindow):
    """Main application window."""
    
    def __init__(self):
        super().__init__()
        
        self.steps: List[Step] = []
        self.current_step_idx: int = 0
        self.session_active: bool = False
        
        # Managers
        self.hotkey_manager = HotkeyManager()
        self.hotkey_manager.capture_triggered.connect(self.on_capture_hotkey)
        self.hotkey_manager.next_triggered.connect(self.on_next_hotkey)
        self.hotkey_manager.prev_triggered.connect(self.on_prev_hotkey)
        self.hotkey_manager.error_occurred.connect(self.show_error)
        
        self._init_ui()
        self._init_default_steps()
        self._update_step_view()
        
        # Register hotkeys
        if not self.hotkey_manager.register():
            QMessageBox.warning(
                self, 
                "Hotkey Registration Failed",
                "Failed to register global hotkeys. The app may require administrator privileges."
            )
        
        logger.info("MainWindow initialized")
    
    def _init_ui(self) -> None:
        """Initialize the user interface."""
        self.setWindowTitle("Test Evidence Helper")
        self.setMinimumSize(700, 550)
        
        # Central widget
        central = QWidget()
        main_layout = QVBoxLayout()
        main_layout.setSpacing(10)
        
        # Document name section
        doc_group = QGroupBox("Document Settings")
        doc_layout = QHBoxLayout()
        doc_layout.addWidget(QLabel("Document Name:"))
        self.doc_name_input = QLineEdit(self._default_doc_name())
        self.doc_name_input.setPlaceholderText("Enter document name...")
        doc_layout.addWidget(self.doc_name_input)
        doc_group.setLayout(doc_layout)
        main_layout.addWidget(doc_group)
        
        # Step navigation section
        step_group = QGroupBox("Current Step")
        step_layout = QVBoxLayout()
        
        self.step_label = QLabel("Step 1/1")
        font = QFont()
        font.setPointSize(12)
        font.setBold(True)
        self.step_label.setFont(font)
        step_layout.addWidget(self.step_label)
        
        self.step_title_label = QLabel("Title:")
        self.step_expected_label = QLabel("Expected:")
        step_layout.addWidget(self.step_title_label)
        step_layout.addWidget(self.step_expected_label)
        
        step_group.setLayout(step_layout)
        main_layout.addWidget(step_group)
        
        # Actual result / notes editor
        editor_group = QGroupBox("Actual Result / Notes")
        editor_layout = QVBoxLayout()
        self.actual_edit = QTextEdit()
        self.actual_edit.setPlaceholderText(
            "Enter actual test result or notes here...\n"
            "This will be saved for the current step.\n\n"
            "You can also paste text/tables from clipboard using Ctrl+Alt+S"
        )
        editor_layout.addWidget(self.actual_edit)
        editor_group.setLayout(editor_layout)
        main_layout.addWidget(editor_group)
        
        # Buttons
        btn_layout_1 = QHBoxLayout()
        
        self.start_btn = QPushButton("▶ Start Session")
        self.start_btn.clicked.connect(self.on_start_session)
        self.start_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold;")
        btn_layout_1.addWidget(self.start_btn)
        
        self.load_excel_btn = QPushButton("📁 Load from Excel File")
        self.load_excel_btn.clicked.connect(self.on_load_excel_file)
        btn_layout_1.addWidget(self.load_excel_btn)
        
        self.paste_excel_btn = QPushButton("📋 Paste from Clipboard")
        self.paste_excel_btn.clicked.connect(self.on_paste_excel_from_clipboard)
        self.paste_excel_btn.setStyleSheet("background-color: #FF9800; color: white; font-weight: bold;")
        btn_layout_1.addWidget(self.paste_excel_btn)
        
        main_layout.addLayout(btn_layout_1)
        
        btn_layout_2 = QHBoxLayout()
        
        self.prev_btn = QPushButton("◀ Previous")
        self.prev_btn.clicked.connect(self.on_prev_step)
        btn_layout_2.addWidget(self.prev_btn)
        
        self.next_btn = QPushButton("Next ▶")
        self.next_btn.clicked.connect(self.on_next_step)
        btn_layout_2.addWidget(self.next_btn)
        
        self.export_btn = QPushButton("💾 Export to DOCX")
        self.export_btn.clicked.connect(self.on_export_docx)
        self.export_btn.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold;")
        btn_layout_2.addWidget(self.export_btn)
        
        main_layout.addLayout(btn_layout_2)
        
        # Hotkey info
        hotkey_info = QLabel(
            f"<b>Global Hotkeys:</b> Capture (Screenshot/Text/Table): <i>{CAPTURE_HOTKEY}</i> | "
            f"Prev: <i>{PREV_STEP_HOTKEY}</i> | Next: <i>{NEXT_STEP_HOTKEY}</i>"
        )
        hotkey_info.setWordWrap(True)
        hotkey_info.setStyleSheet("background-color: #FFF9C4; padding: 5px; border-radius: 3px;")
        main_layout.addWidget(hotkey_info)
        
        central.setLayout(main_layout)
        self.setCentralWidget(central)
        
        # Status bar
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.update_status("Ready - Paste Excel data with '📋 Paste from Clipboard' or load file")
        
        # Auto-save on focus change
        self.actual_edit.focusOutEvent = self._on_editor_focus_out
    
    def _default_doc_name(self) -> str:
        """Generate default document name with timestamp."""
        return f"TestRun_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    def _init_default_steps(self) -> None:
        """Initialize with one empty step."""
        self.steps = [Step(index=1)]
        self.current_step_idx = 0
        logger.debug("Initialized with default empty step")
    
    def update_status(self, message: str, timeout: int = 0) -> None:
        """Update status bar message."""
        self.status_bar.showMessage(message, timeout)
        logger.debug(f"Status: {message}")
    
    def show_error(self, message: str) -> None:
        """Show error message dialog."""
        QMessageBox.critical(self, "Error", message)
        logger.error(f"Error shown to user: {message}")
    
    def show_info(self, title: str, message: str) -> None:
        """Show info message dialog."""
        QMessageBox.information(self, title, message)
    
    # ==================== SESSION MANAGEMENT ====================
    def on_start_session(self) -> None:
        """Start a test session."""
        try:
            self.session_active = True
            self.update_status("Session running - Hotkeys active")
            self.start_btn.setEnabled(False)
            self.start_btn.setText("✓ Session Active")
            logger.info("Test session started")
        except Exception as e:
            logger.error(f"Error starting session: {e}", exc_info=True)
            self.show_error(f"Failed to start session: {e}")
    
    # ==================== STEP NAVIGATION ====================
    def _save_current_step_actual(self) -> None:
        """Save the current actual/notes text to current step."""
        try:
            if 0 <= self.current_step_idx < len(self.steps):
                self.steps[self.current_step_idx].actual = self.actual_edit.toPlainText()
                logger.debug(f"Saved actual text for step {self.current_step_idx + 1}")
        except Exception as e:
            logger.error(f"Error saving step actual: {e}", exc_info=True)
    
    def on_next_step(self) -> None:
        """Navigate to next step (create if at end)."""
        try:
            self._save_current_step_actual()
            
            if self.current_step_idx < len(self.steps) - 1:
                self.current_step_idx += 1
            else:
                # Create new step
                new_step = Step(index=len(self.steps) + 1)
                self.steps.append(new_step)
                self.current_step_idx = len(self.steps) - 1
                logger.info(f"Created new step {new_step.index}")
            
            self._update_step_view()
        except Exception as e:
            logger.error(f"Error navigating to next step: {e}", exc_info=True)
            self.show_error(f"Failed to navigate: {e}")
    
    def on_prev_step(self) -> None:
        """Navigate to previous step."""
        try:
            self._save_current_step_actual()
            
            if self.current_step_idx > 0:
                self.current_step_idx -= 1
                self._update_step_view()
            else:
                self.update_status("Already at first step", 2000)
        except Exception as e:
            logger.error(f"Error navigating to previous step: {e}", exc_info=True)
            self.show_error(f"Failed to navigate: {e}")
    
    def _update_step_view(self) -> None:
        """Update UI to reflect current step."""
        try:
            if not self.steps:
                return
            
            step = self.steps[self.current_step_idx]
            self.step_label.setText(f"Step {step.index}/{len(self.steps)}")
            self.step_title_label.setText(f"<b>Title:</b> {step.title or '<i>No title</i>'}")
            self.step_expected_label.setText(f"<b>Expected:</b> {step.expected or '<i>Not specified</i>'}")
            self.actual_edit.setPlainText(step.actual)
            
            status = (f"Step {step.index} | "
                     f"Screenshots: {len(step.screenshots)} | "
                     f"Notes: {len(step.notes)}")
            self.update_status(status)
            
            logger.debug(f"Updated view for step {step.index}")
        except Exception as e:
            logger.error(f"Error updating step view: {e}", exc_info=True)
    
    def _on_editor_focus_out(self, event) -> None:
        """Handle focus leaving the editor."""
        try:
            self._save_current_step_actual()
        except Exception as e:
            logger.error(f"Error in focus out handler: {e}", exc_info=True)
        finally:
            QTextEdit.focusOutEvent(self.actual_edit, event)
    
    # ==================== HOTKEY HANDLERS ====================
    def on_capture_hotkey(self) -> None:
        """Handle capture hotkey press."""
        if not self.session_active:
            return
        
        try:
            self.capture_clipboard_into_step()
        except Exception as e:
            logger.error(f"Error in capture hotkey handler: {e}", exc_info=True)
            self.show_error(f"Capture failed: {e}")
    
    def on_next_hotkey(self) -> None:
        """Handle next step hotkey."""
        if not self.session_active:
            return
        
        try:
            self.on_next_step()
        except Exception as e:
            logger.error(f"Error in next hotkey handler: {e}", exc_info=True)
    
    def on_prev_hotkey(self) -> None:
        """Handle previous step hotkey."""
        if not self.session_active:
            return
        
        try:
            self.on_prev_step()
        except Exception as e:
            logger.error(f"Error in prev hotkey handler: {e}", exc_info=True)
    
    # ==================== CLIPBOARD CAPTURE ====================
    def capture_clipboard_into_step(self) -> None:
        """Capture clipboard content (image, text, or table) into current step."""
        if not self.steps:
            self.update_status("No steps available", 2000)
            return
        
        step = self.steps[self.current_step_idx]
        
        # Try image first
        try:
            img = ImageGrab.grabclipboard()
        except Exception as e:
            logger.warning(f"Failed to grab clipboard image: {e}")
            img = None
        
        if img is not None:
            # Save screenshot
            try:
                tmp_dir = tempfile.gettempdir()
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S%f')
                filename = f"testhelper_step{step.index}_{timestamp}.png"
                filepath = os.path.join(tmp_dir, filename)
                
                # Validate image
                if not isinstance(img, Image.Image):
                    logger.error(f"Invalid image type: {type(img)}")
                    self.update_status("Invalid image in clipboard", 3000)
                    return
                
                img.save(filepath, "PNG")
                
                if step.add_screenshot(filepath):
                    self.update_status(f"✓ Screenshot captured for step {step.index}", 3000)
                    self._update_step_view()
                else:
                    self.update_status("Failed to add screenshot", 3000)
                
                return
            except Exception as e:
                logger.error(f"Failed to save screenshot: {e}", exc_info=True)
                self.show_error(f"Screenshot save failed: {e}")
                return
        
        # Try Excel/table data (TSV/CSV)
        try:
            df = ClipboardParser.parse_excel_from_clipboard()
            if df is not None and not df.empty:
                # Format as table text
                table_text = df.to_string(index=False)
                step.add_note(f"[Table Data]\n{table_text}")
                
                # Also append to actual field
                existing = step.actual
                if existing:
                    step.actual = f"{existing}\n\n[Pasted Table]\n{table_text}"
                else:
                    step.actual = f"[Pasted Table]\n{table_text}"
                
                self._update_step_view()
                self.update_status(f"✓ Table data captured ({len(df)} rows) for step {step.index}", 3000)
                return
        except Exception as e:
            logger.debug(f"Not Excel/table data: {e}")
        
        # Try plain text
        try:
            clipboard = QApplication.clipboard()
            text = clipboard.text()
            
            if text and text.strip():
                step.add_note(text)
                
                # Also append to actual field
                existing = step.actual
                if existing:
                    step.actual = f"{existing}\n{text}"
                else:
                    step.actual = text
                
                self._update_step_view()
                self.update_status(f"✓ Text captured for step {step.index}", 3000)
            else:
                self.update_status("Clipboard is empty", 2000)
        except Exception as e:
            logger.error(f"Failed to read clipboard text: {e}", exc_info=True)
            self.show_error(f"Clipboard read failed: {e}")
    
    # ==================== EXCEL IMPORT FROM FILE ====================
    def on_load_excel_file(self) -> None:
        """Load test steps from an Excel file."""
        try:
            filepath, _ = QFileDialog.getOpenFileName(
                self,
                "Select Excel File",
                "",
                "Excel Files (*.xlsx *.xls);;All Files (*)"
            )
            
            if not filepath:
                return
            
            logger.info(f"Loading steps from Excel: {filepath}")
            
            # Read Excel
            try:
                df = pd.read_excel(filepath)
            except Exception as e:
                logger.error(f"Failed to read Excel: {e}", exc_info=True)
                self.show_error(f"Failed to read Excel file:\n{e}")
                return
            
            if df.empty:
                self.show_error("Excel file is empty")
                return
            
            # Process with column mapping
            self._process_excel_dataframe(df)
            
        except Exception as e:
            logger.error(f"Error loading Excel: {e}", exc_info=True)
            self.show_error(f"Failed to load Excel:\n{e}")
    
    # ==================== EXCEL PASTE FROM CLIPBOARD ====================
    def on_paste_excel_from_clipboard(self) -> None:
        """Parse Excel data from clipboard (Jira Xray style)."""
        try:
            logger.info("Attempting to parse Excel from clipboard")
            
            df = ClipboardParser.parse_excel_from_clipboard()
            
            if df is None:
                self.show_error(
                    "Could not parse clipboard as Excel data.\n\n"
                    "Please copy data from Excel, Google Sheets, or Jira Xray:\n"
                    "1. Select cells in Excel/Sheets/Xray\n"
                    "2. Copy (Ctrl+C)\n"
                    "3. Click 'Paste from Clipboard' button"
                )
                return
            
            if df.empty:
                self.show_error("Clipboard contains empty data")
                return
            
            logger.info(f"Parsed {len(df)} rows with {len(df.columns)} columns from clipboard")
            
            # Process with column mapping
            self._process_excel_dataframe(df)
            
        except Exception as e:
            logger.error(f"Error pasting Excel from clipboard: {e}", exc_info=True)
            self.show_error(f"Failed to parse clipboard:\n{e}")
    
    def _process_excel_dataframe(self, df: pd.DataFrame) -> None:
        """Process Excel dataframe with column mapping dialog."""
        try:
            # Auto-detect columns
            auto_mapping = ClipboardParser.detect_columns(df)
            
            # Show mapping dialog
            dialog = ColumnMappingDialog(df, auto_mapping, self)
            if dialog.exec_() != QDialog.Accepted:
                return
            
            mapping = dialog.get_mapping()
            
            title_col = mapping['title']
            expected_col = mapping['expected']
            actual_col = mapping['actual']
            skip_empty = mapping['skip_empty']
            
            if not title_col:
                self.show_error("Title column is required")
                return
            
            # Parse rows
            new_steps = []
            for idx, row in df.iterrows():
                try:
                    title_val = str(row[title_col]) if not pd.isna(row[title_col]) else ""
                    expected_val = ""
                    actual_val = ""
                    
                    if expected_col and not pd.isna(row[expected_col]):
                        expected_val = str(row[expected_col])
                    
                    if actual_col and not pd.isna(row[actual_col]):
                        actual_val = str(row[actual_col])
                    
                    # Skip empty titles if option is set
                    if skip_empty and (not title_val.strip() or title_val.strip().lower() in ['nan', 'none', '']):
                        continue
                    
                    step = Step(
                        index=len(new_steps) + 1,
                        title=title_val.strip(),
                        expected=expected_val.strip(),
                        actual=actual_val.strip()
                    )
                    new_steps.append(step)
                except Exception as row_error:
                    logger.warning(f"Skipping row {idx}: {row_error}")
                    continue
            
            if not new_steps:
                self.show_error("No valid steps found in data")
                return
            
            self.steps = new_steps
            self.current_step_idx = 0
            self._update_step_view()
            
            self.show_info(
                "Steps Loaded",
                f"Successfully loaded {len(self.steps)} test steps.\n\n"
                f"Click 'Start Session' to begin capturing evidence.\n"
                f"Use hotkeys to navigate and capture."
            )
            logger.info(f"Loaded {len(self.steps)} steps from Excel data")
            
        except Exception as e:
            logger.error(f"Error processing Excel dataframe: {e}", exc_info=True)
            self.show_error(f"Failed to process data:\n{e}")
    
    # ==================== DOCX EXPORT ====================
    def on_export_docx(self) -> None:
        """Export test results to Word document."""
        try:
            self._save_current_step_actual()
            
            # Validate
            if not self.steps:
                self.show_error("No steps to export")
                return
            
            # Check if all steps are empty
            if all(step.is_empty() for step in self.steps):
                reply = QMessageBox.question(
                    self,
                    "Empty Steps",
                    "All steps appear to be empty. Export anyway?",
                    QMessageBox.Yes | QMessageBox.No
                )
                if reply == QMessageBox.No:
                    return
            
            # Get save path
            suggested_name = self.doc_name_input.text().strip() or self._default_doc_name()
            if not suggested_name.lower().endswith('.docx'):
                suggested_name += '.docx'
            
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Word Document",
                suggested_name,
                "Word Document (*.docx);;All Files (*)"
            )
            
            if not save_path:
                return
            
            if not save_path.lower().endswith('.docx'):
                save_path += '.docx'
            
            # Export
            self.update_status("Exporting document...")
            QApplication.processEvents()  # Update UI
            
            success = DocxExporter.export(
                self.steps,
                save_path,
                title=Path(save_path).stem
            )
            
            if success:
                self.show_info(
                    "Export Successful",
                    f"Document saved to:\n{save_path}"
                )
                self.update_status(f"✓ Exported to {Path(save_path).name}")
            else:
                self.show_error("Failed to export document. Check logs for details.")
                self.update_status("Export failed")
                
        except Exception as e:
            logger.error(f"Error exporting document: {e}", exc_info=True)
            self.show_error(f"Export failed:\n{e}")
            self.update_status("Export failed")
    
    # ==================== CLEANUP ====================
    def closeEvent(self, event) -> None:
        """Handle window close event."""
        try:
            self.hotkey_manager.unregister()
            logger.info("Application closing")
        except Exception as e:
            logger.error(f"Error during cleanup: {e}", exc_info=True)
        finally:
            event.accept()


# ==================== APPLICATION ENTRY POINT ====================
def main() -> int:
    """Main application entry point."""
    try:
        logger.info("=" * 60)
        logger.info("Test Evidence Helper Starting")
        logger.info(f"Python: {sys.version}")
        logger.info(f"Platform: {sys.platform}")
        logger.info("=" * 60)
        
        
        app = QApplication(sys.argv)
        appNamestr = "Test Evidence Helper"
        if platform.system().lower() != 'windws':
            appNamestr = " This is not supported for " + platform.system().lower() + "yet " 
        app.setApplicationName(appNamestr)
        app.setOrganizationName("TestTools")
       
        
        window = MainWindow()
        window.show()
        
        return app.exec_()
        
    except Exception as e:
        logger.critical(f"Fatal error: {e}", exc_info=True)
        print(f"Fatal error: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
